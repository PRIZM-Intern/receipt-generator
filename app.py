
import io
import re
import tempfile
import subprocess
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
import streamlit as st
from PIL import Image, ImageChops, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


# =========================================================
# 원본 5월 영수증 Word 양식 기준
# =========================================================
# 원본 구조:
# - 상단 "영 수 증"은 본문이 아니라 Header
# - Header 제목: 맑은 고딕 22pt Bold, 중앙
# - 본문: 영수증 이미지 1개, 중앙 정렬
# - 이미지 아래: 빈 오른쪽 정렬 문단 1개 + 하단 문구 오른쪽 정렬 1개
# - 하단 문구는 한 문단 안에서 줄바꿈:
#   총 ( n )장 중 ( x )장
#   부서: (부서명) / 한글 본명: (이름)
# - 페이지 여백:
#   top 1701 twips ≈ 3.00cm
#   left/right/bottom 1440 twips ≈ 2.54cm
#   header 851 twips ≈ 1.50cm
#   footer 992 twips ≈ 1.75cm
# =========================================================

FONT_NAME = "맑은 고딕"
BODY_FONT_SIZE_PT = 11
HEADER_FONT_SIZE_PT = 22

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7

TOP_MARGIN_CM = 3.00
BOTTOM_MARGIN_CM = 2.54
LEFT_MARGIN_CM = 2.54
RIGHT_MARGIN_CM = 2.54
HEADER_DISTANCE_CM = 1.50
FOOTER_DISTANCE_CM = 1.75

# 원본 카드전표 이미지가 대체로 폭 9.0~9.7cm, 높이 17.1~19.0cm 수준.
# 가장 긴 택시/세로 영수증도 20.1cm 내외라 아래 max box에 맞춤.
MAX_IMAGE_WIDTH_CM = 9.70
MAX_IMAGE_HEIGHT_CM = 20.10

TITLE_TEXT = "영 수 증"
MAX_UPLOAD_SIZE_MB = 20
MAX_FILE_COUNT = 50
MAX_TOTAL_UPLOAD_SIZE_MB = 100


@dataclass
class ReceiptItem:
    source_name: str
    kind: str
    dt: datetime
    amount: int
    image: Image.Image


def parse_datetime_from_text(text: str):
    patterns = [
        r"(20\d{2})[-./](\d{2})[-./](\d{2})\s+(\d{2}):(\d{2}):(\d{2})",
        r"(20\d{2})[-./](\d{2})[-./](\d{2})\s+(\d{2}):(\d{2})",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            y, mo, d, h, mi, *rest = m.groups()
            s = rest[0] if rest else "00"
            return datetime(int(y), int(mo), int(d), int(h), int(mi), int(s))
    return None


def parse_amount_from_text(text: str):
    m = re.search(r"합계\s*([0-9,]+)\s*원", text)
    if not m:
        m = re.search(r"결제\s*금액\s*([0-9,]+)\s*원", text)
    if m:
        return int(m.group(1).replace(",", ""))
    return None


def sanitize_filename_component(value: str, fallback: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", str(value))
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ._")
    return (cleaned or fallback)[:60]


def crop_white_margin(img: Image.Image, pad_px: int = 24) -> Image.Image:
    """
    PDF 전체 페이지 여백 제거.
    실제 영수증 영역만 자르되 이미지 자체 비율 왜곡은 없음.
    """
    img = img.convert("RGB")
    bg = Image.new("RGB", img.size, (255, 255, 255))
    diff = ImageChops.difference(img, bg).convert("L")
    diff = diff.point(lambda x: 255 if x > 10 else 0)
    bbox = diff.getbbox()

    if not bbox:
        return img

    left, top, right, bottom = bbox
    left = max(0, left - pad_px)
    top = max(0, top - pad_px)
    right = min(img.width, right + pad_px)
    bottom = min(img.height, bottom + pad_px)

    return img.crop((left, top, right, bottom))


def pdf_to_receipt_items(uploaded_file, default_pdf_meal_amount: int):
    data = uploaded_file.getvalue()
    items = []

    with fitz.open(stream=data, filetype="pdf") as doc:
        for page_index in range(len(doc)):
            page = doc[page_index]
            text = page.get_text("text") or ""

            dt = parse_datetime_from_text(text)
            actual_amount = parse_amount_from_text(text)

            if dt is None:
                dt = datetime(2099, 12, 31, 23, 59, 59)

            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            img = crop_white_margin(img)

            name = uploaded_file.name if len(doc) == 1 else f"{uploaded_file.name} - p{page_index + 1}"

            items.append({
                "source_name": name,
                "kind": "PDF 식대",
                "auto_datetime": dt,
                "actual_amount_detected": actual_amount,
                "amount": default_pdf_meal_amount,
                "image": img,
            })

    return items


def image_to_item(uploaded_file):
    img = ImageOps.exif_transpose(Image.open(uploaded_file)).convert("RGB")
    return {
        "source_name": uploaded_file.name,
        "kind": "JPG/기타",
        "auto_datetime": datetime(2099, 12, 31, 23, 59, 59),
        "actual_amount_detected": None,
        "amount": 0,
        "image": img,
    }


def set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_doc_defaults(doc: Document):
    for style_name in ["Normal", "Body Text"]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            style._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
            style.font.size = Pt(BODY_FONT_SIZE_PT)
            style.paragraph_format.space_after = Pt(8)
        except Exception:
            pass


def setup_header(doc: Document):
    header = doc.sections[0].header

    # python-docx 기본 빈 문단 사용
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(TITLE_TEXT)
    set_run_font(run, size_pt=HEADER_FONT_SIZE_PT, bold=True)


def add_picture_centered(doc: Document, image: Image.Image, tmp_dir: Path):
    image_path = tmp_dir / f"receipt_{len(list(tmp_dir.glob('receipt_*.png'))) + 1}.png"
    image.save(image_path)

    w_px, h_px = image.size
    aspect = w_px / h_px

    width_cm = MAX_IMAGE_WIDTH_CM
    height_cm = width_cm / aspect

    if height_cm > MAX_IMAGE_HEIGHT_CM:
        height_cm = MAX_IMAGE_HEIGHT_CM
        width_cm = height_cm * aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))


def add_original_style_footer_text(doc: Document, total: int, idx: int, department: str, person_name: str):
    """
    원본처럼 이미지 아래에 빈 오른쪽 정렬 문단 하나,
    그 다음 문단에 장수/부서 문구를 오른쪽 정렬로 넣음.
    """
    blank = doc.add_paragraph()
    blank.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r1 = p.add_run(f"총 ( {total} )장 중 ( {idx} )장")
    set_run_font(r1, BODY_FONT_SIZE_PT, False)

    p.add_run().add_break()

    r2 = p.add_run(f"부서: ({department}) / 한글 본명: ({person_name})")
    set_run_font(r2, BODY_FONT_SIZE_PT, False)


def create_docx(receipts, month_label: str, department: str, person_name: str, out_path: Path):
    doc = Document()
    set_doc_defaults(doc)

    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(TOP_MARGIN_CM)
    section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
    section.left_margin = Cm(LEFT_MARGIN_CM)
    section.right_margin = Cm(RIGHT_MARGIN_CM)
    section.header_distance = Cm(HEADER_DISTANCE_CM)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)

    setup_header(doc)

    total = len(receipts)

    with tempfile.TemporaryDirectory() as td:
        tmp_dir = Path(td)

        for i, item in enumerate(receipts, start=1):
            if i > 1:
                doc.add_page_break()

            add_picture_centered(doc, item.image, tmp_dir)
            add_original_style_footer_text(doc, total, i, department, person_name)

        # 전체 run 재점검
        for p in doc.paragraphs:
            for r in p.runs:
                # 그림 run은 rPr이 없을 수 있으므로 텍스트 run 위주로 적용
                if r.text:
                    set_run_font(r, BODY_FONT_SIZE_PT, False)

        # header run 재점검
        for p in doc.sections[0].header.paragraphs:
            for r in p.runs:
                if r.text == TITLE_TEXT:
                    set_run_font(r, HEADER_FONT_SIZE_PT, True)

        doc.save(out_path)


def convert_docx_to_pdf(docx_path: Path, out_dir: Path):
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
        )
        pdf_path = out_dir / (docx_path.stem + ".pdf")
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None


st.set_page_config(page_title="월별 영수증 기안서 생성기", layout="wide")

st.title("월별 영수증 기안서 생성기")
st.caption("v3: 원본 5월 Word 양식 기준. Header 제목 22pt Bold, 하단 문구 오른쪽 정렬.")
st.info(
    "업로드한 파일은 문서 생성에만 사용되며 영구 저장하지 않습니다. "
    "생성된 파일은 이 화면에서 바로 다운로드해 주세요."
)

with st.sidebar:
    st.header("기본 정보")
    today = date.today()
    month_label = st.text_input("대상월", value=f"{today.year}년 {today.month}월")
    department = st.text_input("부서명", value="", placeholder="예: 경영지원팀")
    person_name = st.text_input("한글 본명", value="", placeholder="예: 홍길동")
    default_pdf_meal_amount = st.number_input("PDF 식대 인정금액", min_value=0, value=10000, step=1000)

st.subheader("1. 파일 업로드")
uploaded_files = st.file_uploader(
    "PDF/JPG/PNG 영수증을 모두 업로드하세요.",
    type=["pdf", "jpg", "jpeg", "png"],
    accept_multiple_files=True,
)

if uploaded_files:
    if len(uploaded_files) > MAX_FILE_COUNT:
        st.error(f"한 번에 최대 {MAX_FILE_COUNT}개 파일까지 업로드할 수 있습니다.")
        st.stop()

    oversized_files = [
        f.name for f in uploaded_files
        if getattr(f, "size", len(f.getvalue())) > MAX_UPLOAD_SIZE_MB * 1024 * 1024
    ]
    total_upload_size = sum(getattr(f, "size", len(f.getvalue())) for f in uploaded_files)

    if oversized_files:
        st.error(
            f"파일 하나의 최대 크기는 {MAX_UPLOAD_SIZE_MB}MB입니다: "
            + ", ".join(oversized_files)
        )
        st.stop()

    if total_upload_size > MAX_TOTAL_UPLOAD_SIZE_MB * 1024 * 1024:
        st.error(f"전체 업로드 용량은 최대 {MAX_TOTAL_UPLOAD_SIZE_MB}MB입니다.")
        st.stop()

    raw_items = []
    for f in uploaded_files:
        try:
            ext = Path(f.name).suffix.lower()
            if ext == ".pdf":
                raw_items.extend(pdf_to_receipt_items(f, int(default_pdf_meal_amount)))
            else:
                raw_items.append(image_to_item(f))
        except Exception:
            st.error(f"{f.name}: 파일을 읽을 수 없습니다. 손상 여부를 확인해 주세요.")

    if not raw_items:
        st.warning("처리할 수 있는 영수증이 없습니다.")
        st.stop()

    st.subheader("2. 자동 인식 결과 검수")
    st.write("PDF는 거래일시가 자동 인식됩니다. JPG/PNG는 결제일시와 금액을 직접 입력해 주세요.")

    rows = []
    for item in raw_items:
        dt = item["auto_datetime"]
        rows.append({
            "사용": True,
            "파일명": item["source_name"],
            "구분": item["kind"],
            "결제일시": "" if dt.year == 2099 else dt.strftime("%Y-%m-%d %H:%M:%S"),
            "금액": item["amount"],
            "참고_실제금액": item["actual_amount_detected"] if item["actual_amount_detected"] else "",
        })

    df = pd.DataFrame(rows)
    edited = st.data_editor(
        df,
        num_rows="fixed",
        use_container_width=True,
        column_config={
            "파일명": st.column_config.TextColumn("파일명", disabled=True),
            "구분": st.column_config.TextColumn("구분", disabled=True),
            "결제일시": st.column_config.TextColumn("결제일시", help="예: 2026-07-01 13:20:00"),
            "금액": st.column_config.NumberColumn("금액", min_value=0, step=100),
            "참고_실제금액": st.column_config.TextColumn("참고_실제금액", disabled=True),
        },
    )

    st.subheader("3. 이미지 미리보기")
    preview_cols = st.columns(4)
    for idx, item in enumerate(raw_items[:8]):
        with preview_cols[idx % 4]:
            st.image(item["image"], caption=item["source_name"], use_container_width=True)

    if st.button("DOCX/PDF 생성하기", type="primary"):
        final_items = []
        errors = []

        if not department.strip():
            errors.append("부서명을 입력해 주세요.")
        if not person_name.strip():
            errors.append("한글 본명을 입력해 주세요.")

        for idx, row in edited.iterrows():
            if not row["사용"]:
                continue

            try:
                dt = datetime.strptime(str(row["결제일시"]).strip(), "%Y-%m-%d %H:%M:%S")
            except Exception:
                errors.append(f"{row['파일명']}: 결제일시 형식 오류")
                continue

            try:
                amount = int(row["금액"])
            except Exception:
                errors.append(f"{row['파일명']}: 금액 오류")
                continue

            base = raw_items[idx]
            final_items.append(ReceiptItem(
                source_name=base["source_name"],
                kind=base["kind"],
                dt=dt,
                amount=amount,
                image=base["image"],
            ))

        if errors:
            st.error("수정이 필요한 항목이 있습니다.")
            for e in errors:
                st.write("- " + e)
        elif not final_items:
            st.error("사용할 영수증이 없습니다.")
        else:
            final_items.sort(key=lambda x: x.dt)

            safe_month = sanitize_filename_component(month_label, "월별")
            safe_department = sanitize_filename_component(department, "부서")
            safe_person_name = sanitize_filename_component(person_name, "이름")
            output_name = f"{safe_month}_영수증_{safe_department}_{safe_person_name}"

            try:
                with tempfile.TemporaryDirectory() as td:
                    out_dir = Path(td)
                    docx_path = out_dir / f"{output_name}.docx"
                    create_docx(final_items, month_label, department, person_name, docx_path)
                    pdf_path = convert_docx_to_pdf(docx_path, out_dir)
                    docx_data = docx_path.read_bytes()
                    pdf_data = pdf_path.read_bytes() if pdf_path and pdf_path.exists() else None
            except Exception:
                st.error("문서 생성 중 오류가 발생했습니다. 입력 파일을 확인한 뒤 다시 시도해 주세요.")
                st.stop()

            total_amount = sum(x.amount for x in final_items)
            st.success(f"생성 완료: 총 {len(final_items)}장 / 합계 {total_amount:,}원")

            summary = pd.DataFrame([
                {
                    "순서": i + 1,
                    "결제일시": x.dt.strftime("%Y-%m-%d %H:%M:%S"),
                    "구분": x.kind,
                    "파일명": x.source_name,
                    "금액": x.amount,
                }
                for i, x in enumerate(final_items)
            ])
            st.dataframe(summary, use_container_width=True)

            st.download_button(
                "DOCX 다운로드",
                docx_data,
                file_name=f"{output_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            if pdf_data:
                st.download_button(
                    "PDF 다운로드",
                    pdf_data,
                    file_name=f"{output_name}.pdf",
                    mime="application/pdf",
                )
            else:
                st.warning("PDF 변환은 LibreOffice가 설치된 환경에서만 자동 생성됩니다. 현재 환경에서는 DOCX만 다운로드하세요.")
else:
    st.info("먼저 PDF/JPG/PNG 파일을 업로드하세요.")
